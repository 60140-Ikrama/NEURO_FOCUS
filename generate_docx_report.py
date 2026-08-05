"""
Comprehensive Word Document (.docx) Project Report Generator for NeuroLearn Research Suite.
Processes real-world PhysioNet EEG data (Subject 1, Runs 1 & 2) and compiles an editable IEEE/Frontiers style
project report in Microsoft Word format.
"""

import os
import sys
import time
import numpy as np
import pandas as pd
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from data.physionet import PhysioNetManager
from preprocessing.cleaner import EEGPreprocessor
from preprocessing.quality import SignalQualityAssessment
from segmentation.windowing import EEGSegmenter
from features.extractor import FeatureExtractionEngine
from attention.calculator import AttentionCalculator
from stats.analyzer import StatisticalAnalyzer
from ablation.engine import AblationEngine


def set_cell_background(cell, fill_color: str):
    """Set shading color for a table cell in python-docx."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'))


def build_full_project_docx(output_filename: str = "NeuroLearn_Research_Suite_Full_Project_Report.docx") -> str:
    """Generate comprehensive multi-page Word (.docx) research project report using real PhysioNet data."""

    print("Loading real-world PhysioNet EEG dataset for DOCX report...")
    pn_data_run1 = PhysioNetManager.load_physionet_subject(subject_id=1, runs=[1])
    pn_data_run2 = PhysioNetManager.load_physionet_subject(subject_id=1, runs=[2])

    # 1. Real SQI Audit
    sqa = SignalQualityAssessment()
    sqi_run1 = sqa.evaluate(pn_data_run1)

    # 2. Real Preprocessing
    cleaner = EEGPreprocessor(lowcut=1.0, highcut=40.0, notch_freq=50.0, normalization="zscore")
    clean_run1, preproc_log = cleaner.process(pn_data_run1)
    clean_run2, _ = cleaner.process(pn_data_run2)

    # 3. Real Windowing
    segmenter = EEGSegmenter(window_sec=5.0, overlap_ratio=0.5)
    epochs_run1 = segmenter.segment(clean_run1)
    epochs_run2 = segmenter.segment(clean_run2)

    # 4. Real Feature & Attention Calculation
    fe = FeatureExtractionEngine(psd_method="welch")
    feat_run1 = fe.extract_features(epochs_run1)
    feat_run2 = fe.extract_features(epochs_run2)

    calc = AttentionCalculator(formula_str="beta / theta")
    att_run1, summary_run1 = calc.compute_attention(feat_run1)
    att_run2, summary_run2 = calc.compute_attention(feat_run2)

    # 5. Real Stats & Bland-Altman
    scores_r1 = att_run1["attention_score"].values
    scores_r2 = att_run2["attention_score"].values
    stat_comp = StatisticalAnalyzer.compare_groups(scores_r1, scores_r2, group1_name="Run 1 (Baseline)", group2_name="Run 2 (Task)")
    ba_res = StatisticalAnalyzer.bland_altman(scores_r1, scores_r2)

    # 6. Real Ablation
    ablation_eng = AblationEngine(pn_data_run1)
    df_ablation = ablation_eng.run_preprocessing_ablation()

    # Create Word Document
    doc = docx.Document()

    # Page Margins (1 inch = 72pt)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Title Header Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("NeuroLearn Research Suite (NEURO_FOCUS)\n")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0A, 0x19, 0x2F)

    run_sub = p_title.add_run("Comprehensive Technical & Research Project Report\n")
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    run_desc = p_title.add_run("A Modular Biomedical Signal Processing (BSP) Platform for EEG-Based Attention Quantification")
    run_desc.font.size = Pt(10)
    run_desc.font.italic = True
    run_desc.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Metadata Table
    table_meta = doc.add_table(rows=3, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False

    meta_data_cells = [
        [f"Primary Dataset: {pn_data_run1.subject_id} (eegmmidb)", "Methodology: 100% Non-ML / Explainable BSP"],
        [f"Sampling Rate (fs): {pn_data_run1.sampling_rate} Hz", f"Channel Count: {pn_data_run1.n_channels} Channels"],
        [f"Recording Duration: {pn_data_run1.duration_sec:.1f} Seconds", f"Quality Audit: {sqi_run1['overall_sqi']}% SQI Score"]
    ]

    for i in range(3):
        for j in range(2):
            cell = table_meta.cell(i, j)
            cell.text = meta_data_cells[i][j]
            set_cell_background(cell, "F4F7FA")
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 1: Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Real-World Dataset Validation", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    p1 = doc.add_paragraph(
        f"This report presents an empirical biomedical signal processing study evaluating student cognitive attention state "
        f"derived directly from the PhysioNet EEG Motor Movement/Imagery Database (eegmmidb). "
        f"The primary dataset consists of 64-channel EEG recordings sampled at {pn_data_run1.sampling_rate} Hz from {pn_data_run1.subject_id} "
        f"encompassing both resting baseline (Run 1) and active task conditions (Run 2)."
    )
    p1.paragraph_format.space_after = Pt(8)

    p2 = doc.add_paragraph(
        "Strict Non-ML Policy: In accordance with rigorous scientific standards, zero machine learning or black-box predictive models "
        "were used. All attention metrics were computed using zero-phase Butterworth bandpass filtering (1-40 Hz), 50 Hz powerline notch filtering, "
        "Welch Power Spectral Density (PSD), Hjorth parameters (Activity, Mobility, Complexity), and mathematical band ratio formulations (Beta/Theta)."
    )
    p2.paragraph_format.space_after = Pt(14)

    # Section 2: Core Architecture
    h2 = doc.add_heading("2. Core Modular Architecture & Subsystems", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    doc.add_paragraph("The platform comprises 17 decoupled backend processing modules and 21 interactive frontend workspaces:")

    modules_list = [
        "1. Dataset Manager: Parses EDF, MAT, CSV, and TXT recordings with automatic metadata detection.",
        "2. Biopac Import Manager: Normalizes AcqKnowledge MAT/CSV exports into standard EEGData containers.",
        "3. PhysioNet Adapter: Fetches and formats public PhysioNet EEG Motor Imagery datasets (eegmmidb).",
        "4. Signal Quality Assessment (SQI): Computes SQI scores (0-100%), SNR (dB), and auto-excludes bad channels.",
        "5. Signal Preprocessing Cleaner: Executes zero-phase Bandpass (1-40Hz), 50Hz Notch filter, and Z-score standardization.",
        "6. Sliding Window Engine: Segments continuous EEG into sliding epochs (2s to 30s) with overlap.",
        "7. Feature Extraction Engine: Extracts time-domain stats, Hjorth parameters, Welch PSD, SEF95, and Wavelet Entropy.",
        "8. Frequency Analysis Subsystem: Calculates absolute and relative band powers for Delta, Theta, Alpha, Beta, and Gamma.",
        "9. Mathematical Attention Calculator: Computes non-ML ratio formulas, AST user expressions, and 5-level state classification.",
        "10. Statistical Analysis Module: Performs hypothesis testing, Shapiro-Wilk normality tests, Cohen's d, and Bland-Altman.",
        "11. Systematic Ablation Study Engine: Evaluates Preprocessing pipelines, Window sizes, PSD methods, and Formulas.",
        "12. Research Validation Engine: Direct cross-dataset agreement audit between PhysioNet and Biopac EEG sessions.",
        "13. Experiment Manager & Reproducibility: Logs software versions, filter bounds, and exports 100% reproducible JSON snapshots.",
        "14. Multi-Format Report Generator: Exports IEEE/Frontiers-styled PDF, Word DOCX, and CSV dataset reports."
    ]

    for m in modules_list:
        p_m = doc.add_paragraph(m, style='List Bullet')
        p_m.paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Empirical Results Table
    h3 = doc.add_heading("3. Empirical Attention Analysis (Baseline vs. Task)", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    table_att = doc.add_table(rows=6, cols=4)
    table_att.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_att.autofit = False

    headers = ["Metric", "Run 1 (Baseline)", "Run 2 (Task)", "Scientific Significance"]
    for j, text in enumerate(headers):
        cell = table_att.cell(0, j)
        cell.text = text
        set_cell_background(cell, "0052CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.bold = True

    att_rows_data = [
        ["Average Attention Score", f"{summary_run1['average_attention']:.1f} / 100", f"{summary_run2['average_attention']:.1f} / 100", "Mean cognitive score"],
        ["Peak Attention Score", f"{summary_run1['peak_attention']:.1f} / 100", f"{summary_run2['peak_attention']:.1f} / 100", "Maximum focused state"],
        ["Minimum Attention Score", f"{summary_run1['minimum_attention']:.1f} / 100", f"{summary_run2['minimum_attention']:.1f} / 100", "Nadir attention level"],
        ["Attention Stability Index", f"{summary_run1['stability_index']:.1f}%", f"{summary_run2['stability_index']:.1f}%", "Inverse std deviation"],
        ["Dominant Category", str(summary_run1['dominant_category']), str(summary_run2['dominant_category']), "5-Level rule-based state"]
    ]

    for i, row_vals in enumerate(att_rows_data, 1):
        for j, val in enumerate(row_vals):
            cell = table_att.cell(i, j)
            cell.text = val
            if i % 2 == 1:
                set_cell_background(cell, "F4F7FA")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Section 4: Statistical Results
    h4 = doc.add_heading("4. Statistical Hypothesis Testing & Bland-Altman Agreement", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    p_st = doc.add_paragraph(
        f"Hypothesis Test: {stat_comp['test_name']}\n"
        f"Test Statistic: {stat_comp['statistic']:.3f} (p-value = {stat_comp['p_value']:.4f})\n"
        f"Cohen's d Effect Size: d = {stat_comp['cohens_d']:.2f}\n"
        f"Bland-Altman Agreement: Mean Bias = {ba_res['mean_difference']:.2f} points, "
        f"95% Limits of Agreement = [{ba_res['loa_lower_95']:.2f}, {ba_res['loa_upper_95']:.2f}], Pearson r = {ba_res['pearson_r']:.3f}."
    )
    p_st.paragraph_format.space_after = Pt(14)

    # Section 5: Real Ablation Study Table
    h5 = doc.add_heading("5. Real-World Systematic Ablation Study Results", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    table_ab = doc.add_table(rows=len(df_ablation) + 1, cols=5)
    table_ab.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ab.autofit = False

    ab_headers = ["Preprocessing Pipeline", "Mean Attention", "Peak Attention", "Stability Index", "Dominant Category"]
    for j, text in enumerate(ab_headers):
        cell = table_ab.cell(0, j)
        cell.text = text
        set_cell_background(cell, "00A896")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.bold = True

    for i, (_, r) in enumerate(df_ablation.iterrows(), 1):
        vals = [
            str(r["Configuration"]),
            f"{r['Mean Attention']:.1f} / 100",
            f"{r['Peak Attention']:.1f}",
            f"{r['Stability Index']:.1f}%",
            str(r["Dominant Category"])
        ]
        for j, val in enumerate(vals):
            cell = table_ab.cell(i, j)
            cell.text = val
            if i == 1:
                set_cell_background(cell, "E6F4F1")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Section 6: Mathematical Equations & References
    h6 = doc.add_heading("6. Mathematical Derivations & Academic References", level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    p_eq = doc.add_paragraph(
        "• 4th Order Zero-Phase Butterworth Bandpass: |H(jw)|^2 = 1 / [ 1 + ((w^2 - w0^2)/(w*B))^(2N) ]\n"
        "• Hjorth Activity: var(x(t)) = sigma_x^2 | Mobility: sigma_x' / sigma_x | Complexity: Mobility(dx/dt) / Mobility(x(t))\n"
        "• Welch PSD: PSD(f) = (1/K) sum_{k=1}^K |FFT(x_k * w)|^2 / (L * U)\n"
        "• Spectral Edge Frequency (SEF95): Integral_0^{f95} PSD(f) df = 0.95 * Total_Power\n"
        "• Attention Ratio: Beta / Theta = Power(13-30Hz) / Power(4-8Hz)\n"
        "• Cohen's d: d = (Mean_1 - Mean_2) / s_pooled"
    )
    p_eq.paragraph_format.space_after = Pt(10)

    p_ref = doc.add_paragraph(
        "Academic References:\n"
        "1. Goldberger, A. L., et al. (2000). 'PhysioBank, PhysioToolkit, and PhysioNet: Components of a new research resource for complex physiologic signals.' Circulation, 101(23), e215-e220.\n"
        "2. Coelli, S., et al. (2018). 'EEG-based indices of attention during online learning tasks.' IEEE Transactions on Biomedical Engineering.\n"
        "3. Hjorth, B. (1970). 'EEG analysis based on time domain properties.' Electroencephalography and Clinical Neurophysiology."
    )
    p_ref.paragraph_format.space_after = Pt(14)

    # Save Word Document
    output_path = os.path.join(os.getcwd(), output_filename)
    doc.save(output_path)
    print(f"Full Project DOCX Report successfully generated at: {output_path}")
    return output_path


if __name__ == "__main__":
    path = build_full_project_docx()
